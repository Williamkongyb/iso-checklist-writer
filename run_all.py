#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ISO20000/ISO27001 检查表一键生成器
用法: python run_all.py --records-dir "D:/..." --standard ISO20000 --audit-date 2026-02-01 --company-name "山东翰林" [--template "path"] [--skip-read] [--quiet]

三步合一：Step1读取记录 → Step2 DeepSeek生成内容 → Step3格式克隆输出
全程只消耗DeepSeek token，WorkBuddy仅做编排调用此脚本

添加 --quiet 参数：安静模式，只输出最终文件路径，详细日志写入文件
"""

import os, sys, json, time, re, shutil, argparse, subprocess, tempfile
from pathlib import Path
from datetime import datetime, timedelta

# 强制UTF-8（Windows兼容）
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

# === 全局日志控制 ===
QUIET_MODE = False
LOG_FILE = None

def log(msg, level="INFO"):
    """日志输出：quiet模式下只输出到文件，非quiet模式输出到stdout"""
    global QUIET_MODE, LOG_FILE
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_line = f"[{timestamp}] [{level}] {msg}"
    
    # 写入日志文件
    if LOG_FILE:
        try:
            with open(LOG_FILE, 'a', encoding='utf-8') as f:
                f.write(log_line + '\n')
        except:
            pass
    
    # 非quiet模式输出到stdout
    if not QUIET_MODE:
        print(msg, flush=True)

def final_output(msg):
    """最终输出：无论是否quiet模式，都输出到stdout（供用户查看结果）"""
    print(msg, flush=True)
    # 也写入日志
    global LOG_FILE
    if LOG_FILE:
        try:
            with open(LOG_FILE, 'a', encoding='utf-8') as f:
                f.write(msg + '\n')
        except:
            pass

# === DeepSeek API Config ===
API_KEY = os.environ.get("DEEPSEEK_API_KEY", "sk-71569ae60b494983afb1ebc45c8440ab")
BASE_URL = os.environ.get("DEEPSEEK_BASE_URL", "https://api.deepseek.com")
MODEL = "deepseek-chat"
BATCH_SIZE = 6

# === 模板路径 ===
TEMPLATES = {
    "ISO20000": r"D:\合规知识库-compliance-kb\checklists-检查表\CEPREI A-196-D-2024.08 信息技术服务管理体系检查表ISO20000-2018.docx",
    "ISO27001": r"D:\合规知识库-compliance-kb\checklists-检查表\CEPREI A-152-M-2024.08 信息安全管理体系检查表ISO27001-2022.doc",
}
# 输出到脚本所在目录的output子目录（避免沙箱权限问题）
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")

# ============================================================
# Step1: 读取客户记录（递归读取所有文件）
# ============================================================
def read_docx(path, max_chars=8000):
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t: parts.append(t)
        for tb in doc.tables:
            for row in tb.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells: parts.append(" | ".join(cells))
        return "\n".join(parts)[:max_chars]
    except Exception as e:
        return f"[ERROR] {e}"

def read_doc_antiword(path, max_chars=5000):
    try:
        tmp = tempfile.mktemp(suffix=".doc")
        ps = f"Copy-Item -Path '{path}' -Destination '{tmp}' -Force"
        subprocess.run(["powershell", "-Command", ps], capture_output=True, timeout=30)
        r = subprocess.run(["antiword", "-m", "UTF-8", tmp], capture_output=True, timeout=30)
        try: os.remove(tmp)
        except: pass
        if r.returncode != 0: return "[ERROR] antiword failed"
        return r.stdout.decode("utf-8", errors="replace")[:max_chars]
    except Exception as e:
        return f"[ERROR] {e}"

def read_xlsx(path, max_chars=3000):
    try:
        from openpyxl import load_workbook
        wb = load_workbook(path, read_only=True)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(max_row=30, values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells: parts.append(" | ".join(cells))
        wb.close()
        return "\n".join(parts)[:max_chars]
    except Exception as e:
        return f"[ERROR] {e}"

def smart_read(filepath):
    ext = filepath.suffix.lower()
    if ext == ".docx": return read_docx(str(filepath))
    elif ext == ".doc": return read_doc_antiword(str(filepath))
    elif ext in (".xlsx", ".xls"): return read_xlsx(str(filepath))
    elif ext == ".txt":
        try: return filepath.read_text(encoding="utf-8", errors="replace")[:3000]
        except: return "[ERROR]"
    return None

def step1_read_records(records_dir):
    """读取客户记录目录，递归读取所有文件"""
    root = Path(records_dir)
    if not root.exists():
        log(f"ERROR: {root} not found!", "INFO")
        sys.exit(1)
    
    log(f"\n[Step1] Reading records from {root.name}", "INFO")
    
    all_records = {}
    total_chars = 0
    file_count = 0
    
    # 递归读取所有 .doc/.docx/.xlsx/.txt 文件
    for f in sorted(root.rglob("*")):
        if f.is_dir() or f.name.startswith("~$"):
            continue
        if f.suffix.lower() not in (".doc", ".docx", ".xlsx", ".xls", ".txt"):
            continue
        
        rel = f.relative_to(root)
        text = smart_read(f)
        if text and not text.startswith("[ERROR]") and len(text) > 30:
            all_records[str(rel)] = text
            total_chars += len(text)
            file_count += 1
    
    # 合并输出
    merged = ""
    for path, text in all_records.items():
        merged += f"\n{'='*60}\nFILE: {path}\n{'='*60}\n{text}\n"
    
    log(f"  Files: {file_count}, Chars: {total_chars}", "INFO")
    return merged

# ============================================================
# Step2: DeepSeek API 生成检查表内容
# ============================================================
def call_deepseek(messages, max_tokens=8192, max_retries=3):
    import urllib.request
    data = json.dumps({
        "model": MODEL,
        "messages": messages,
        "max_tokens": max_tokens,
        "temperature": 0.1
    }).encode("utf-8")
    
    req = urllib.request.Request(
        f"{BASE_URL}/chat/completions",
        data=data,
        headers={"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"}
    )
    
    for attempt in range(max_retries):
        try:
            resp = urllib.request.urlopen(req, timeout=120)
            result = json.loads(resp.read().decode("utf-8"))
            if "choices" in result:
                return result["choices"][0]["message"]["content"], result.get("usage", {})
            else:
                err = json.dumps(result, ensure_ascii=False)[:200]
                log(f"  API error: {err}", "INFO")
                time.sleep(3)
        except Exception as e:
            log(f"  Retry {attempt+1}: {e}", "INFO")
            time.sleep(5 * (attempt + 1))
    return "", {}

def parse_template_clauses(template_path, standard):
    """解析模板，提取需要填写的条款"""
    from docx import Document
    doc = Document(template_path)
    content_col = 3  # 内容列index
    seen = set()
    clauses = []
    
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if len(cells) <= content_col: continue
            key = (t_idx, r_idx, content_col)
            if key in seen: continue
            
            content_text = cells[content_col]
            req = cells[0] if cells[0] and len(cells[0]) > 2 else (cells[1] if len(cells) > 1 and cells[1] else "")
            
            if not content_text and req and len(req) > 2:
                seen.add(key)
                clauses.append({
                    "t": t_idx, "r": r_idx, "c": content_col,
                    "req": req[:300],
                    "preview": " | ".join(cells[:4])[:200]
                })
    return clauses

RULES_PROMPT = """你是ISO合规审查检查表内容生成专家。严格按以下规则：

1. 只填写"内容"列，不修改其他
2. 内容格式="[制度文件名]：[3-5句简要摘抄]。查[记录证据]..."
3. 制度文件名必须来自客户记录，严禁编造
4. 话术模板：
   - 程序文件："制定《XXX管理程序》，对[管理对象]的[流程/活动]进行了规范。明确了[职责分工/流程步骤/输出要求]。"
   - 服务管理："制定《XXX服务管理手册/程序》，规划并建立了[管理体系/流程]。明确了[服务目标/绩效指标]。"
   - 过程管理："制定《XXX过程管理程序》，对[过程名称]的[输入/活动/输出]进行了规范。"
5. 审核话术："查《XXX》，[摘要]"；结尾"验证无异常"/"符合要求"
6. 合同/项目用前-中-后话术结构
7. 三条线对齐仅在综合部分"抽查项目"
8. 内审管评时间线核查
9. 五大维度：时间、人员、地点、发现、结果
10. 时间锚点校验（审核日前1年）
11. 文件名/编号必须来自实际记录
12. 发现问题优先于写检查表
14. 涉及时间必须写日期，至少到月份"""

def step2_deepseek_generate(records_text, template_path, standard, company_name, audit_date):
    """调用DeepSeek API批量生成检查表内容"""
    log(f"\n[Step2] DeepSeek content generation", "INFO")
    log(f"  Company: {company_name}", "INFO")
    log(f"  Audit date: {audit_date}", "INFO")
    log(f"  Standard: {standard}", "INFO")
    
    # 截断记录文本
    if len(records_text) > 25000:
        records_text = records_text[:25000] + "\n...[truncated]"
        log(f"  Records truncated to 25000 chars", "INFO")
    else:
        log(f"  Records: {len(records_text)} chars", "INFO")
    
    # 解析模板条款
    clauses = parse_template_clauses(template_path, standard)
    log(f"  Template clauses: {len(clauses)}", "INFO")
    
    # 分批
    batches = [clauses[i:i+BATCH_SIZE] for i in range(0, len(clauses), BATCH_SIZE)]
    log(f"  Batches: {len(batches)} (size={BATCH_SIZE})", "INFO")
    
    window_start = (datetime.strptime(audit_date, "%Y-%m-%d") - timedelta(days=365)).strftime("%Y-%m-%d")
    all_content = {}
    total_tokens = 0
    issues_found = []
    
    for b_idx, batch in enumerate(batches):
        log(f"\n  --- Batch {b_idx+1}/{len(batches)} ({len(batch)} clauses) ---", "INFO")
        
        clauses_desc = ""
        for i, c in enumerate(batch):
            clauses_desc += f"\n<<CLAUSE_{i}>>\nTable={c['t']} Row={c['r']} Col={c['c']}\nRequirement: {c['req']}\n<</CLAUSE_{i}>>\n"
        
        prompt = f"""请为以下{len(batch)}个ISO{standard}检查条款生成"内容列"审核记录。

企业：{company_name}
审核日期：{audit_date}
有效窗口：{window_start} 至 {audit_date}

客户记录：
{records_text}

待填写条款：
{clauses_desc}

输出格式（必须严格遵守）：
<<ANSWER_0>>
内容列文字...
<</ANSWER_0>>
<<ANSWER_1>>
内容列文字...
<</ANSWER_1>>
（以此类推）

规则：
1. 按规则2-6话术组织文字
2. 文件名/编号/人名/日期必须来自客户记录，严禁编造
3. 必须包含时间（至少到月份）
4. 缺少的信息写"（待确认）"
5. 只输出内容列文字"""
        
        messages = [
            {"role": "system", "content": RULES_PROMPT},
            {"role": "user", "content": prompt}
        ]
        
        content, usage = call_deepseek(messages, max_tokens=8192)
        batch_tokens = usage.get("total_tokens", 0)
        total_tokens += batch_tokens
        log(f"  Tokens: {batch_tokens}", "INFO")
        
        # 解析输出
        for i, c in enumerate(batch):
            pattern = f'<<ANSWER_{i}>>(.*?)<</ANSWER_{i}>>'
            m = re.search(pattern, content, re.DOTALL)
            if m:
                text = m.group(1).strip()
                key = f"{c['t']}|{c['r']}|{c['c']}"
                all_content[key] = text
                log(f"    T{c['t']}R{c['r']}: {text[:60]}...", "INFO")
                # 检测问题
                if "待确认" in text or "待补充" in text:
                    issues_found.append({"clause": c['req'][:50], "issue": text[:200]})
            else:
                key = f"{c['t']}|{c['r']}|{c['c']}"
                all_content[key] = "（待补充）"
                log(f"    T{c['t']}R{c['r']}: PARSE_FAILED", "INFO")
        
        # 限流
        if b_idx < len(batches) - 1:
            time.sleep(1)
    
    log(f"\n  DeepSeek tokens: {total_tokens}", "INFO")
    return all_content, total_tokens, issues_found

# ============================================================
# Step 3: 格式克隆 - 填充模板输出.docx
# ============================================================
def step3_format_clone(content_dict, template_path, company_name, standard, audit_date):
    """格式克隆：复制模板 → 填充内容 → 保存"""
    log(f"\n[Step3] Format clone", "INFO")
    log(f"  Content entries: {len(content_dict)}", "INFO")
    
    # 生成输出文件名
    now = datetime.now()
    period = "上午" if now.hour < 12 else "下午" if now.hour < 18 else "晚上"
    date_str = now.strftime("%Y-%m-%d")
    timestamp = now.strftime("%H%M%S")
    output_name = f"{date_str}_{period}_{company_name}_{standard}检查表_DS{timestamp}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_name)
    
    # 确保输出目录存在
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    # 复制模板
    shutil.copy2(template_path, output_path)
    log(f"  Template copied", "INFO")
    
    # 填充
    from docx import Document
    doc = Document(output_path)
    filled = 0
    failed = 0
    
    for key, text in content_dict.items():
        try:
            parts = key.split("|")
            t_idx, r_idx, c_idx = int(parts[0]), int(parts[1]), int(parts[2])
        except:
            continue
        
        try:
            table = doc.tables[t_idx]
            row = table.rows[r_idx]
            cell = row.cells[c_idx]
            
            # 清空现有内容
            for p in cell.paragraphs:
                for run in p.runs:
                    run.text = ""
            
            # 添加新内容
            if cell.paragraphs:
                cell.paragraphs[0].add_run(text)
            else:
                cell.add_paragraph(text)
            filled += 1
        except Exception as e:
            failed += 1
            if failed <= 5:
                log(f"  FILL ERROR T{t_idx}R{r_idx}C{c_idx}: {e}", "INFO")
    
    doc.save(output_path)
    
    log(f"  Filled: {filled}", "INFO")
    log(f"  Failed: {failed}", "INFO")
    final_output(f"✅ 检查表已生成：{output_path}")
    return output_path, filled, failed

# ============================================================
# Step 4: 生成问题清单
# ============================================================
def step4_generate_issues(issues_found, company_name, standard, audit_date, records_text):
    """调用DeepSeek生成问题清单和整改建议"""
    if not issues_found:
        log(f"\n[Step4] No issues found, skipping issue list generation", "INFO")
        return None
    
    log(f"\n[Step4] Generating issue list ({len(issues_found)} potential issues)", "INFO")
    
    issues_desc = "\n".join([f"- 条款: {i['clause']}\n  问题: {i['issue']}" for i in issues_found])
    
    prompt = f"""基于以下审核发现，生成正式的问题清单和整改建议。

企业：{company_name}
标准：{standard}
审核日期：{audit_date}

发现的潜在问题：
{issues_desc}

客户记录摘要（前5000字符）：
{records_text[:5000]}

请输出：
1. 问题清单（编号、条款、问题描述、严重程度、涉及记录）
2. 整改建议（针对每个问题给出具体整改措施、责任人建议、完成时限建议）

注意：只列出有证据支持的问题，不要编造问题。"""
    
    messages = [
        {"role": "system", "content": "你是ISO合规审查专家，擅长发现体系运行中的问题并给出整改建议。"},
        {"role": "user", "content": prompt}
    ]
    
    content, usage = call_deepseek(messages, max_tokens=4096)
    if not content:
        log(f"  DeepSeek returned empty", "INFO")
        return None
    
    # 保存为文本（后续可转为docx）
    now = datetime.now()
    period = "上午" if now.hour < 12 else "下午" if now.hour < 18 else "晚上"
    date_str = now.strftime("%Y-%m-%d")
    issues_name = f"{date_str}_{period}_{company_name}_{standard}问题清单.txt"
    issues_path = os.path.join(OUTPUT_DIR, issues_name)
    
    with open(issues_path, "w", encoding="utf-8") as f:
        f.write(content)
    
    log(f"  Tokens: {usage.get('total_tokens', 0)}", "INFO")
    final_output(f"✅ 问题清单已生成：{issues_path}")
    return issues_path

# ============================================================
# Main
# ============================================================
def main():
    parser = argparse.ArgumentParser(description="ISO20000/ISO27001 检查表一键生成")
    parser.add_argument("--records-dir", required=True, help="客户记录目录路径")
    parser.add_argument("--standard", required=True, choices=["ISO20000", "ISO27001"], help="标准类型")
    parser.add_argument("--audit-date", required=True, help="审核日期 YYYY-MM-DD")
    parser.add_argument("--company-name", required=True, help="公司名称")
    parser.add_argument("--template", help="模板文件路径（默认自动选择）")
    parser.add_argument("--skip-read", action="store_true", help="跳过Step1（使用已有records.txt）")
    parser.add_argument("--records-file", help="已有的记录文本文件路径")
    parser.add_argument("--quiet", action="store_true", help="安静模式：只输出最终结果，详细日志写入文件")
    args = parser.parse_args()
    
    # 设置全局quiet模式
    global QUIET_MODE, LOG_FILE
    QUIET_MODE = args.quiet
    
    if QUIET_MODE:
        # 创建日志文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        log_filename = f"log_{timestamp}.txt"
        LOG_FILE = os.path.join(OUTPUT_DIR, log_filename)
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        with open(LOG_FILE, 'w', encoding='utf-8') as f:
            f.write(f"ISO检查表生成日志 - {datetime.now()}\n")
            f.write(f"=" * 80 + "\n\n")
        log(f"安静模式启动，日志文件: {LOG_FILE}", "INFO")
    
    log("=" * 60, "INFO")
    log(f"ISO Checklist Generator (DeepSeek-driven)", "INFO")
    log(f"Company: {args.company_name}", "INFO")
    log(f"Standard: {args.standard}", "INFO")
    log(f"Audit date: {args.audit_date}", "INFO")
    log(f"Quiet mode: {QUIET_MODE}", "INFO")
    log("=" * 60, "INFO")
    
    # 确定模板
    template_path = args.template or TEMPLATES.get(args.standard)
    if not template_path or not os.path.exists(template_path):
        log(f"ERROR: Template not found for {args.standard}: {template_path}", "INFO")
        sys.exit(1)
    log(f"Template: {os.path.basename(template_path)}", "INFO")
    
    total_start = time.time()
    
    # Step 1: 读取记录
    if args.skip_read and args.records_file and os.path.exists(args.records_file):
        with open(args.records_file, "r", encoding="utf-8") as f:
            records_text = f.read()
        log(f"[Step1] SKIPPED (using existing file, {len(records_text)} chars)", "INFO")
    elif args.skip_read:
        # 尝试默认路径
        default_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "hanlin_records.txt")
        if os.path.exists(default_path):
            with open(default_path, "r", encoding="utf-8") as f:
                records_text = f.read()
            log(f"[Step1] SKIPPED (using default file, {len(records_text)} chars)", "INFO")
        else:
            log(f"ERROR: --skip-read but no records file found", "INFO")
            sys.exit(1)
    else:
        records_text = step1_read_records(args.records_dir)
    
    # Step 2: DeepSeek生成
    content_dict, ds_tokens, issues = step2_deepseek_generate(
        records_text, template_path, args.standard, args.company_name, args.audit_date
    )
    
    # Step 3: 格式克隆
    checklist_path, filled, failed = step3_format_clone(
        content_dict, template_path, args.company_name, args.standard, args.audit_date
    )
    
    # Step 4: 问题清单
    issues_path = step4_generate_issues(
        issues, args.company_name, args.standard, args.audit_date, records_text
    )
    
    # 汇总
    total_time = time.time() - total_start
    log(f"\n{'='*60}", "INFO")
    final_output(f"✅ 完成！耗时 {total_time:.1f}s")
    final_output(f"  检查表：{checklist_path}")
    final_output(f"  问题清单：{issues_path or '无'}")
    final_output(f"  DeepSeek tokens：{ds_tokens}")
    final_output(f"  填充：{filled}/{filled+failed}")
    log(f"{'='*60}", "INFO")

if __name__ == "__main__":
    main()
